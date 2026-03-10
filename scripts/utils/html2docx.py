#!/usr/bin/env python3
"""
HTML to DOCX converter for WeChat Official Account
简化版 - 专注于内容转换
"""

import sys
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import NavigableString

def html_to_docx(html_path: str, output_path: str = None):
    """Convert HTML file to DOCX format"""
    
    if output_path is None:
        output_path = Path(html_path).with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    # Read HTML content
    with open(html_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    soup = BeautifulSoup(content, 'html.parser')
    doc = Document()
    
    # Set default font for Chinese support
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(14)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Process elements
    for element in soup.body.children:
        process_element(element, doc)
    
    # Save document
    doc.save(output_path)
    print(f"✅ 转换成功：{output_path}")
    return output_path

def process_element(element, doc):
    """递归处理 HTML 元素"""
    if not hasattr(element, 'name'):
        return
    
    tag_name = element.name
    
    # H1 - 主标题
    if tag_name == 'h1':
        p = doc.add_heading(element.get_text(strip=True), level=1)
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(22)
        run.font.bold = True
    
    # H2 - 二级标题
    elif tag_name == 'h2':
        text = element.get_text(strip=True)
        p = doc.add_paragraph()
        run = p.add_run('▍' + text)
        run.font.size = Pt(17)
        run.font.bold = True
    
    # H3 - 三级标题
    elif tag_name == 'h3':
        p = doc.add_heading(element.get_text(strip=True), level=3)
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(16)
        run.font.bold = True
    
    # P - 段落
    elif tag_name == 'p':
        parent = element.parent
        is_special = False
        
        if parent and parent.has_attr('class'):
            classes = ' '.join(parent.get('class', []))
            
            if 'intro-box' in classes:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.5)
                run = p.add_run('【导语】\n')
                run.font.bold = True
                add_formatted_text(element, p)
                return
            elif 'key-points' in classes:
                p = doc.add_paragraph()
                run = p.add_run('💡 核心要点：\n')
                run.font.bold = True
                add_formatted_text(element, p)
                return
            elif 'highlight-box' in classes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('✨ 金句 ✨\n\n')
                run.font.bold = True
                add_formatted_text(element, p)
                return
            elif 'interaction' in classes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('\n📬 互动话题\n\n')
                run.font.bold = True
                run.font.size = Pt(16)
                add_formatted_text(element, p)
                return
            elif 'footer' in classes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_formatted_text(element, p, color=RGBColor(153, 153, 153), size=Pt(13))
                return
        
        if not is_special:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_text(element, p)
    
    # UL/OL - 列表
    elif tag_name == 'ul' or tag_name == 'ol':
        is_ordered = tag_name == 'ol'
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet' if not is_ordered else 'List Number')
            add_formatted_text(li, p)
    
    # BLOCKQUOTE - 引用
    elif tag_name == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run('" ')
        run.font.size = Pt(16)
        add_formatted_text(element, p, color=RGBColor(85, 85, 85), size=Pt(15))
        run = p.add_run(' "')
        run.font.size = Pt(16)
    
    # TABLE - 表格
    elif tag_name == 'table':
        rows = element.find_all('tr')
        if rows:
            cols = len(rows[0].find_all(['th', 'td']))
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            
            for row_idx, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for col_idx, cell in enumerate(cells):
                    if row_idx < len(rows) and col_idx < cols:
                        table_cell = table.cell(row_idx, col_idx)
                        table_cell.text = cell.get_text(strip=True)
                        
                        # 表头加粗
                        if row_idx == 0 or cell.name == 'th':
                            for paragraph in table_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(14)
    
    # DIV - 处理特殊 div
    elif tag_name == 'div':
        if element.has_attr('class'):
            classes = ' '.join(element.get('class', []))
            
            # 分隔线
            if 'divider' in classes:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(16)
                run = p.add_run('─' * 50)
                run.font.color.rgb = RGBColor(234, 234, 234)
                run.font.size = Pt(8)
            
            # CODE-BLOCK - 代码块
            elif 'code-block' in classes:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.5)
                run = p.add_run(element.get_text(strip=True))
                run.font.name = 'Courier New'
                run.font.size = Pt(13)
                run.font.bold = True

def add_formatted_text(element, paragraph, color=None, size=Pt(14)):
    """添加格式化的文本"""
    if color is None:
        color = RGBColor(51, 51, 51)
    
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            text = ' '.join(text.split())
            run = paragraph.add_run(text)
            run.font.color.rgb = color
            run.font.size = size
        return
    
    for child in element.children:
        if isinstance(child, NavigableString):
            add_formatted_text(child, paragraph, color, size)
        elif hasattr(child, 'name'):
            if child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.font.bold = True
                run.font.color.rgb = color
                run.font.size = size
            elif child.name == 'br':
                paragraph.add_run('\n')
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.font.italic = True
                run.font.color.rgb = color
                run.font.size = size
            else:
                add_formatted_text(child, paragraph, color, size)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法：python html2docx.py <html 文件> [输出文件]")
        sys.exit(1)
    
    html_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    html_to_docx(html_file, output_file)
