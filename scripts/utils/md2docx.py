#!/usr/bin/env python3
"""
Markdown to DOCX converter for patent disclosure documents
支持中文、表格、标题层级、列表等格式
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def markdown_to_docx(md_path: str, output_path: str = None):
    """Convert markdown file to DOCX format"""
    
    if output_path is None:
        output_path = Path(md_path).with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font for Chinese support
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Process content line by line
    lines = content.split('\n')
    current_list = []
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty horizontal rules
        if stripped.startswith('---') or stripped.startswith('***'):
            i += 1
            continue
        
        # Handle tables (markdown table format)
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Parse table row
            cells = [cell.strip() for cell in stripped.split('|')]
            # Remove empty first/last cells from markdown table syntax
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            # Skip separator row (contains ---)
            if not any('---' in cell for cell in cells):
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table - create it
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if row_idx < len(table_rows) and col_idx < len(row_data):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_data
                            # Bold first row (header)
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
            in_table = False
            table_rows = []
        
        # Handle headers
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            doc.add_heading(stripped[2:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[3:], level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[4:], level=4)
        elif stripped.startswith('##### '):
            doc.add_heading(stripped[5:], level=5)
        # Handle bold text
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True
        # Handle lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            item = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)
        elif stripped.startswith('1. ') or re.match(r'^\d+\.\s', stripped):
            item = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.add_run(item)
        # Handle regular paragraphs
        elif stripped:
            # Handle inline bold
            text = stripped
            if '**' in text:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(text)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_data in enumerate(row_data):
                if row_idx < len(table_rows) and col_idx < len(row_data):
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_data
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    # Save document
    doc.save(output_path)
    print(f"✅ 转换成功：{output_path}")
    return output_path

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法：python md2docx.py <markdown 文件> [输出文件]")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    markdown_to_docx(md_file, output_file)
